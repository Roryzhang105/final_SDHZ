#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
insert_delivery_receipt.py

向《送达回证》Word 模板自动填充：
  1. “送达文书 名称及文号”（仿宋_GB2312 12pt，水平居中）
  2. 送达人、送达时间、送达地点、受送达人（同字体/字号，可选）
  3. 将两张图片分别插入：
       • “备注”右侧单元格
       • 最底下一整栏
依赖：python-docx  pillow
"""

from pathlib import Path
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


# ────────────────── 辅助函数 ──────────────────
def add_centered_picture(paragraph, img_path: Path, max_width_inch: float):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    with Image.open(img_path) as im:
        w_px, h_px = im.size
        dpi = im.info.get("dpi", (96, 96))[0]

    w_in, h_in = w_px / dpi, h_px / dpi
    if w_in > max_width_inch:
        scale = max_width_inch / w_in
        w_in, h_in = w_in * scale, h_in * scale

    run.add_picture(str(img_path), width=Inches(w_in), height=Inches(h_in))


def write_centered_text(cell, text: str):
    """清空并写入居中文本：仿宋_GB2312 12pt"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "仿宋_GB2312"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    run.font.size = Pt(12)


def fill_cell_by_label(table, label: str, value: str):
    """在表格中找到包含 label 的单元格，把 value 写到其右侧格"""
    print(f"DEBUG: 查找标签 '{label}', 要填充的值: '{value}'")
    
    if value is None:
        print(f"DEBUG: 值为None，跳过填充标签 '{label}'")
        return  # 未提供则跳过

    # 先打印所有单元格的内容，帮助调试
    print(f"DEBUG: 表格中所有单元格内容:")
    for row_idx, row in enumerate(table.rows):
        for cell_idx, cell in enumerate(row.cells):
            cell_text = cell.text.strip()
            if cell_text:  # 只打印非空单元格
                print(f"  - 行{row_idx}列{cell_idx}: '{cell_text}'")

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if label in cell.text.strip():
                # 优先选右侧格；若已是行尾则写自身
                target = row.cells[idx + 1] if idx + 1 < len(row.cells) else cell
                print(f"DEBUG: 找到标签 '{label}' 在单元格 '{cell.text.strip()}'，将在目标单元格填充: '{value}'")
                write_centered_text(target, value)
                return
    
    print(f"DEBUG: 警告 - 未找到包含'{label}'的单元格")
    # 改为警告而不是抛出异常，避免因为找不到某个标签就中断整个流程
    # raise RuntimeError(f"未找到包含"{label}"的单元格")


# ────────────────── 主逻辑 ──────────────────
def process_document(
    template_doc: Path,
    output_doc: Path,
    doc_title: str,
    pic_note: Path,
    pic_footer: Path,
    sender: str | None = None,
    send_time: str | None = None,
    send_location: str | None = None,
    receiver: str | None = None,
    max_width_inch: float = 2.8,
):
    doc = Document(template_doc)
    if not doc.tables:
        raise ValueError("模板中未发现表格")
    table = doc.tables[0]

    # 1) 送达文书 名称及文号
    # 改进的标题填充逻辑，支持分行情况
    title_filled = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell_text = cell.text.replace('\n', '').replace(' ', '')  # 移除换行和空格
            # 检查是否包含"送达文书"关键词（支持分行情况）
            if "送达文书" in cell_text and ("名称" in cell_text or "文号" in cell_text):
                # 找到包含相关关键词的单元格，在其右侧填充标题
                if idx + 1 < len(row.cells):
                    write_centered_text(row.cells[idx + 1], doc_title)
                    print(f"DEBUG: 已在'{cell.text.strip()}'右侧填充文档标题: '{doc_title}'")
                    title_filled = True
                    break
        if title_filled:
            break
    
    # 如果上述方法未成功，尝试更宽松的匹配
    if not title_filled:
        print("DEBUG: 尝试宽松匹配模式")
        for row in table.rows:
            for idx, cell in enumerate(row.cells):
                if "送达文书" in cell.text:
                    # 只要包含"送达文书"就尝试填充
                    if idx + 1 < len(row.cells):
                        write_centered_text(row.cells[idx + 1], doc_title)
                        print(f"DEBUG: 宽松匹配 - 已在'{cell.text.strip()}'右侧填充文档标题: '{doc_title}'")
                        title_filled = True
                        break
            if title_filled:
                break
    
    if not title_filled:
        print(f"WARNING: 未能找到合适的位置填充文档标题: '{doc_title}'")

    # 2) 其它可选字段
    fill_cell_by_label(table, "送达人", sender)
    fill_cell_by_label(table, "送达时间", send_time)
    fill_cell_by_label(table, "送达地点", send_location)
    fill_cell_by_label(table, "受送达人", receiver)

    # 3) 图片
    fill_cell_by_label(table, "备注", "")  # 定位
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "备注" == cell.text.strip():
                note_target = row.cells[c_idx + 1]
                note_target.text = ""
                add_centered_picture(note_target.paragraphs[0], pic_note, max_width_inch)
                break

    footer_cell = table.rows[-1].cells[0]
    footer_cell.text = ""
    add_centered_picture(footer_cell.paragraphs[0], pic_footer, max_width_inch)

    doc.save(output_doc)
    print(f"✓ 已生成：{output_doc}")


# ────────────────── CLI ──────────────────
def main():
    ap = argparse.ArgumentParser(description="批量自动填写《送达回证》")
    ap.add_argument("--template", required=True, help="模板 .docx")
    ap.add_argument("--output", required=True, help="输出 .docx")
    ap.add_argument("--doc-title", required=True, help="送达文书 名称及文号")
    ap.add_argument("--note-img", required=True, help="插入备注栏右侧图片")
    ap.add_argument("--footer-img", required=True, help="插入最底栏图片")
    # 新增可选字段
    ap.add_argument("--sender", help="送达人")
    ap.add_argument("--send-time", help="送达时间")
    ap.add_argument("--send-location", help="送达地点")
    ap.add_argument("--receiver", help="受送达人")
    ap.add_argument("--max-width", type=float, default=2.8, help="图片最大宽度(inch)")

    args = ap.parse_args()
    
    # 添加参数接收日志
    print(f"DEBUG: Word生成脚本接收到的参数:")
    print(f"  - template: {args.template}")
    print(f"  - output: {args.output}")
    print(f"  - doc_title: '{args.doc_title}'")
    print(f"  - sender: '{args.sender}'")
    print(f"  - send_time: '{args.send_time}'")
    print(f"  - send_location: '{args.send_location}'")
    print(f"  - receiver: '{args.receiver}'")
    print(f"  - note_img: {args.note_img}")
    print(f"  - footer_img: {args.footer_img}")

    process_document(
        template_doc=Path(args.template),
        output_doc=Path(args.output),
        doc_title=args.doc_title,
        pic_note=Path(args.note_img),
        pic_footer=Path(args.footer_img),
        sender=args.sender,
        send_time=args.send_time,
        send_location=args.send_location,
        receiver=args.receiver,
        max_width_inch=args.max_width,
    )


if __name__ == "__main__":
    main()

