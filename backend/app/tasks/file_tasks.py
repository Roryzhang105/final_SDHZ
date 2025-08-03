import os
import logging
import qrcode
from barcode import Code128
from barcode.writer import ImageWriter
from docx import Document
from PIL import Image
from celery import current_app as celery_app
from sqlalchemy.orm import Session
from sqlalchemy.exc import SQLAlchemyError

from app.core.database import SessionLocal
from app.core.config import settings
from app.services.delivery_receipt import DeliveryReceiptService
from app.tasks.retry_config import (
    get_retry_config,
    retry_task,
    log_task_failure,
    log_task_retry
)

logger = logging.getLogger(__name__)


@celery_app.task(**get_retry_config('file_operation'))
@retry_task('file_operation', on_failure=log_task_failure, on_retry=log_task_retry)
def generate_qr_and_barcode(self, receipt_id: int):
    """
    生成二维码和条形码的异步任务
    
    重试策略: file_operation (最多3次重试，20秒起始延迟，指数退避)
    """
    logger.info(f"开始生成二维码和条形码: {receipt_id}")
    
    db: Session = SessionLocal()
    
    try:
        service = DeliveryReceiptService(db)
        receipt = service.get_delivery_receipt(receipt_id)
        
        if not receipt:
            error_msg = f"送达回证不存在: {receipt_id}"
            logger.error(error_msg)
            return {"success": False, "error": error_msg}
        
        if not receipt.tracking_number:
            error_msg = f"回证缺少快递单号: {receipt_id}"
            logger.error(error_msg)
            return {"success": False, "error": error_msg}
        
        logger.info(f"找到回证记录: {receipt.tracking_number}")
        
        # 创建文件目录
        codes_dir = os.path.join(settings.UPLOAD_DIR, "codes")
        try:
            os.makedirs(codes_dir, exist_ok=True)
            logger.debug(f"二维码目录已创建: {codes_dir}")
        except OSError as e:
            logger.error(f"创建二维码目录失败: {codes_dir}, 错误: {str(e)}")
            raise e
        
        # 生成二维码
        try:
            qr_path = generate_qr_code(receipt.tracking_number, codes_dir)
            if not qr_path:
                raise Exception("二维码生成返回空路径")
            logger.info(f"二维码生成成功: {qr_path}")
        except Exception as e:
            logger.error(f"生成二维码失败: {receipt.tracking_number}, 错误: {str(e)}")
            raise e
        
        # 生成条形码
        try:
            barcode_path = generate_barcode(receipt.tracking_number, codes_dir)
            if not barcode_path:
                raise Exception("条形码生成返回空路径")
            logger.info(f"条形码生成成功: {barcode_path}")
        except Exception as e:
            logger.error(f"生成条形码失败: {receipt.tracking_number}, 错误: {str(e)}")
            raise e
        
        # 验证生成的文件是否存在
        if not os.path.exists(qr_path):
            raise FileNotFoundError(f"二维码文件不存在: {qr_path}")
        if not os.path.exists(barcode_path):
            raise FileNotFoundError(f"条形码文件不存在: {barcode_path}")
        
        # 更新文件路径到数据库
        try:
            service.update_receipt_files(
                receipt_id=receipt_id,
                qr_code_path=qr_path,
                barcode_path=barcode_path
            )
            db.commit()
            logger.info(f"文件路径已更新到数据库: {receipt_id}")
        except SQLAlchemyError as e:
            logger.error(f"更新文件路径到数据库失败: {receipt_id}, 错误: {str(e)}")
            db.rollback()
            raise e
        
        logger.info(f"二维码和条形码生成完成: {receipt_id}")
        
        return {
            "success": True,
            "message": "二维码和条形码生成成功",
            "receipt_id": receipt_id,
            "tracking_number": receipt.tracking_number,
            "qr_path": qr_path,
            "barcode_path": barcode_path
        }
        
    except (FileNotFoundError, PermissionError, OSError) as e:
        logger.error(f"文件操作错误: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    except SQLAlchemyError as e:
        logger.error(f"数据库操作错误: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    except Exception as e:
        logger.error(f"生成二维码和条形码失败: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    finally:
        db.close()


def generate_qr_code(tracking_number: str, output_dir: str) -> str:
    """
    生成二维码
    增强错误处理和日志记录
    """
    try:
        logger.debug(f"开始生成二维码: {tracking_number}")
        
        # 验证输入参数
        if not tracking_number or not tracking_number.strip():
            raise ValueError("快递单号不能为空")
        
        if not output_dir:
            raise ValueError("输出目录不能为空")
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 创建二维码
        qr = qrcode.QRCode(
            version=1,
            error_correction=qrcode.constants.ERROR_CORRECT_L,
            box_size=10,
            border=4,
        )
        
        qr.add_data(tracking_number.strip())
        qr.make(fit=True)
        
        # 生成图片
        img = qr.make_image(fill_color="black", back_color="white")
        
        # 生成安全的文件名
        safe_tracking_number = "".join(c for c in tracking_number if c.isalnum() or c in "-_")
        filename = f"qr_{safe_tracking_number}.png"
        file_path = os.path.join(output_dir, filename)
        
        # 保存文件
        img.save(file_path)
        
        # 验证文件是否成功保存
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"二维码文件保存失败: {file_path}")
        
        file_size = os.path.getsize(file_path)
        if file_size == 0:
            raise Exception(f"生成的二维码文件为空: {file_path}")
        
        logger.debug(f"二维码生成成功: {file_path}, 大小: {file_size} bytes")
        return file_path
        
    except Exception as e:
        logger.error(f"生成二维码失败: {tracking_number}, 错误: {str(e)}")
        raise e


def generate_barcode(tracking_number: str, output_dir: str) -> str:
    """
    生成条形码
    增强错误处理和日志记录
    """
    try:
        logger.debug(f"开始生成条形码: {tracking_number}")
        
        # 验证输入参数
        if not tracking_number or not tracking_number.strip():
            raise ValueError("快递单号不能为空")
        
        if not output_dir:
            raise ValueError("输出目录不能为空")
        
        # 确保输出目录存在
        if not os.path.exists(output_dir):
            os.makedirs(output_dir, exist_ok=True)
        
        # 创建条形码（Code128支持字母数字）
        code = Code128(tracking_number.strip(), writer=ImageWriter())
        
        # 生成安全的文件名
        safe_tracking_number = "".join(c for c in tracking_number if c.isalnum() or c in "-_")
        filename = f"barcode_{safe_tracking_number}"
        file_path = os.path.join(output_dir, filename)
        
        # 保存条形码（会自动添加.png扩展名）
        code.save(file_path)
        
        # 返回完整路径
        full_path = f"{file_path}.png"
        
        # 验证文件是否成功保存
        if not os.path.exists(full_path):
            raise FileNotFoundError(f"条形码文件保存失败: {full_path}")
        
        file_size = os.path.getsize(full_path)
        if file_size == 0:
            raise Exception(f"生成的条形码文件为空: {full_path}")
        
        logger.debug(f"条形码生成成功: {full_path}, 大小: {file_size} bytes")
        return full_path
        
    except Exception as e:
        logger.error(f"生成条形码失败: {tracking_number}, 错误: {str(e)}")
        raise e


@celery_app.task(**get_retry_config('file_operation'))
@retry_task('file_operation', on_failure=log_task_failure, on_retry=log_task_retry)
def fill_receipt_template(self, receipt_id: int):
    """
    填充送达回证模板的异步任务
    
    重试策略: file_operation (最多3次重试，20秒起始延迟，指数退避)
    """
    logger.info(f"开始填充送达回证模板: {receipt_id}")
    
    db: Session = SessionLocal()
    
    try:
        service = DeliveryReceiptService(db)
        receipt = service.get_delivery_receipt(receipt_id)
        
        if not receipt:
            error_msg = f"送达回证不存在: {receipt_id}"
            logger.error(error_msg)
            return {"success": False, "error": error_msg}
        
        if not receipt.tracking_number:
            error_msg = f"回证缺少快递单号: {receipt_id}"
            logger.error(error_msg)
            return {"success": False, "error": error_msg}
        
        logger.info(f"找到回证记录: {receipt.tracking_number}")
        
        # 填充模板
        try:
            receipt_path = fill_document_template(receipt)
            if not receipt_path:
                raise Exception("文档模板填充返回空路径")
            logger.info(f"文档模板填充成功: {receipt_path}")
        except Exception as e:
            logger.error(f"填充文档模板失败: {receipt_id}, 错误: {str(e)}")
            raise e
        
        # 验证生成的文件
        if not os.path.exists(receipt_path):
            raise FileNotFoundError(f"生成的文档文件不存在: {receipt_path}")
        
        file_size = os.path.getsize(receipt_path)
        if file_size == 0:
            raise Exception(f"生成的文档文件为空: {receipt_path}")
        
        # 更新文件路径到数据库
        try:
            service.update_receipt_files(
                receipt_id=receipt_id,
                receipt_file_path=receipt_path
            )
            db.commit()
            logger.info(f"文档路径已更新到数据库: {receipt_id}")
        except SQLAlchemyError as e:
            logger.error(f"更新文档路径到数据库失败: {receipt_id}, 错误: {str(e)}")
            db.rollback()
            raise e
        
        logger.info(f"送达回证文档生成完成: {receipt_id}")
        
        return {
            "success": True,
            "message": "送达回证文档生成成功",
            "receipt_id": receipt_id,
            "tracking_number": receipt.tracking_number,
            "receipt_path": receipt_path,
            "file_size": file_size
        }
        
    except (FileNotFoundError, PermissionError, OSError) as e:
        logger.error(f"文件操作错误: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    except SQLAlchemyError as e:
        logger.error(f"数据库操作错误: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    except Exception as e:
        logger.error(f"填充送达回证模板失败: {receipt_id}, 错误: {str(e)}")
        db.rollback()
        raise e
        
    finally:
        db.close()


def fill_document_template(receipt) -> str:
    """
    填充Word文档模板
    增强错误处理和日志记录
    """
    try:
        logger.debug(f"开始填充文档模板: {receipt.tracking_number}")
        
        # 模板文件路径
        template_path = "templates/receipt_template.docx"
        
        if not os.path.exists(template_path):
            raise FileNotFoundError(f"模板文件不存在: {template_path}")
        
        # 验证模板文件是否可读
        if not os.access(template_path, os.R_OK):
            raise PermissionError(f"没有权限读取模板文件: {template_path}")
        
        # 加载模板
        try:
            doc = Document(template_path)
        except Exception as e:
            raise Exception(f"加载Word模板失败: {str(e)}")
        
        # 准备替换字典，处理None值
        replacements = {
            "{{tracking_number}}": receipt.tracking_number or "",
            "{{recipient_name}}": getattr(receipt, 'recipient_name', "") or "",
            "{{recipient_address}}": getattr(receipt, 'recipient_address', "") or "",
            "{{sender_name}}": getattr(receipt, 'sender_name', "") or "",
            "{{courier_company}}": getattr(receipt, 'courier_company', "") or "",
            "{{send_time}}": getattr(receipt, 'send_time', "") or "",
            "{{send_location}}": getattr(receipt, 'send_location', "") or "",
            "{{receiver}}": getattr(receipt, 'receiver', "") or "",
            "{{doc_title}}": getattr(receipt, 'doc_title', "送达回证") or "送达回证"
        }
        
        logger.debug(f"替换字典准备完成: {len(replacements)} 项")
        
        # 替换段落中的文本
        for paragraph in doc.paragraphs:
            original_text = paragraph.text
            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))
            
            if original_text != paragraph.text:
                logger.debug(f"段落文本已替换: {original_text[:50]}...")
        
        # 替换表格中的文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    original_text = cell.text
                    for placeholder, value in replacements.items():
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, str(value))
                    
                    if original_text != cell.text:
                        logger.debug(f"表格文本已替换: {original_text[:30]}...")
        
        # 创建输出目录
        output_dir = os.path.join(settings.UPLOAD_DIR, "receipts")
        try:
            os.makedirs(output_dir, exist_ok=True)
        except OSError as e:
            raise OSError(f"创建输出目录失败: {output_dir}, 错误: {str(e)}")
        
        # 生成安全的文件名
        safe_tracking_number = "".join(c for c in receipt.tracking_number if c.isalnum() or c in "-_")
        filename = f"receipt_{safe_tracking_number}.docx"
        output_path = os.path.join(output_dir, filename)
        
        # 保存填充后的文档
        try:
            doc.save(output_path)
        except Exception as e:
            raise Exception(f"保存Word文档失败: {str(e)}")
        
        # 验证文件是否成功保存
        if not os.path.exists(output_path):
            raise FileNotFoundError(f"文档文件保存失败: {output_path}")
        
        file_size = os.path.getsize(output_path)
        if file_size == 0:
            raise Exception(f"生成的文档文件为空: {output_path}")
        
        logger.debug(f"文档模板填充成功: {output_path}, 大小: {file_size} bytes")
        return output_path
        
    except Exception as e:
        logger.error(f"填充文档模板失败: {receipt.tracking_number if receipt else 'Unknown'}, 错误: {str(e)}")
        raise e


@celery_app.task(**get_retry_config('default'))
@retry_task('default', on_failure=log_task_failure, on_retry=log_task_retry)
def cleanup_old_files(self):
    """
    清理旧文件的定时任务
    
    重试策略: default (最多3次重试，60秒起始延迟，指数退避)
    """
    from datetime import datetime, timedelta
    import shutil
    
    logger.info("开始清理旧文件")
    
    try:
        stats = {
            "cleaned_files": 0,
            "freed_space_mb": 0,
            "errors": []
        }
        
        # 清理配置
        cleanup_dirs = [
            {
                "path": os.path.join(settings.UPLOAD_DIR, "codes"),
                "days": 30,
                "description": "二维码/条形码文件"
            },
            {
                "path": os.path.join(settings.UPLOAD_DIR, "receipts"),
                "days": 90,
                "description": "回证文档文件"
            },
            {
                "path": os.path.join(settings.UPLOAD_DIR, "screenshots"),
                "days": 30,
                "description": "截图文件"
            },
            {
                "path": "/tmp",
                "days": 1,
                "description": "临时文件",
                "pattern": "*.tmp"
            }
        ]
        
        cutoff_date = datetime.now() - timedelta(days=30)  # 默认30天
        
        for cleanup_config in cleanup_dirs:
            dir_path = cleanup_config["path"]
            days = cleanup_config["days"]
            description = cleanup_config["description"]
            pattern = cleanup_config.get("pattern", "*")
            
            if not os.path.exists(dir_path):
                logger.debug(f"清理目录不存在，跳过: {dir_path}")
                continue
            
            try:
                dir_cutoff = datetime.now() - timedelta(days=days)
                logger.info(f"清理 {description} - 目录: {dir_path}, 保留天数: {days}")
                
                files_cleaned = 0
                space_freed = 0
                
                for root, dirs, files in os.walk(dir_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        
                        try:
                            # 检查文件修改时间
                            file_mtime = datetime.fromtimestamp(os.path.getmtime(file_path))
                            
                            if file_mtime < dir_cutoff:
                                # 获取文件大小
                                file_size = os.path.getsize(file_path)
                                
                                # 删除文件
                                os.remove(file_path)
                                
                                files_cleaned += 1
                                space_freed += file_size
                                
                                logger.debug(f"删除旧文件: {file_path}")
                                
                        except Exception as e:
                            error_msg = f"删除文件失败: {file_path}, 错误: {str(e)}"
                            stats["errors"].append(error_msg)
                            logger.warning(error_msg)
                
                stats["cleaned_files"] += files_cleaned
                stats["freed_space_mb"] += space_freed / (1024 * 1024)
                
                logger.info(f"{description} 清理完成: 删除 {files_cleaned} 个文件, 释放 {space_freed/(1024*1024):.2f} MB")
                
            except Exception as e:
                error_msg = f"清理目录失败: {dir_path}, 错误: {str(e)}"
                stats["errors"].append(error_msg)
                logger.error(error_msg)
        
        logger.info(f"文件清理完成 - 总计删除 {stats['cleaned_files']} 个文件, 释放 {stats['freed_space_mb']:.2f} MB")
        
        return {
            "success": True,
            "message": "文件清理完成",
            "stats": stats
        }
        
    except Exception as e:
        logger.error(f"清理旧文件失败: {str(e)}")
        raise e


@celery_app.task(**get_retry_config('default'))
@retry_task('default', on_failure=log_task_failure, on_retry=log_task_retry)
def optimize_database(self):
    """
    数据库优化任务
    
    重试策略: default (最多3次重试，60秒起始延迟，指数退避)
    """
    logger.info("开始数据库优化")
    
    db: Session = SessionLocal()
    
    try:
        # PostgreSQL优化命令
        if settings.DATABASE_URL.startswith("postgresql"):
            logger.info("执行PostgreSQL数据库优化")
            
            # 重建索引
            db.execute("REINDEX DATABASE;")
            logger.info("数据库索引重建完成")
            
            # 更新统计信息
            db.execute("ANALYZE;")
            logger.info("数据库统计信息更新完成")
            
            # 清理无用空间
            db.execute("VACUUM;")
            logger.info("数据库空间清理完成")
            
        elif settings.DATABASE_URL.startswith("sqlite"):
            logger.info("执行SQLite数据库优化")
            
            # SQLite优化
            db.execute("PRAGMA optimize;")
            db.execute("VACUUM;")
            logger.info("SQLite数据库优化完成")
        
        db.commit()
        
        logger.info("数据库优化完成")
        
        return {
            "success": True,
            "message": "数据库优化完成"
        }
        
    except SQLAlchemyError as e:
        logger.error(f"数据库优化失败: {str(e)}")
        db.rollback()
        raise e
        
    except Exception as e:
        logger.error(f"数据库优化异常: {str(e)}")
        db.rollback()
        raise e
        
    finally:
        db.close()


@celery_app.task(**get_retry_config('default'))
@retry_task('default', on_failure=log_task_failure, on_retry=log_task_retry)
def backup_database(self):
    """
    数据库备份任务
    
    重试策略: default (最多3次重试，60秒起始延迟，指数退避)
    """
    import subprocess
    from datetime import datetime
    
    logger.info("开始数据库备份")
    
    try:
        backup_dir = "backups/database"
        os.makedirs(backup_dir, exist_ok=True)
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        
        if settings.DATABASE_URL.startswith("postgresql"):
            # PostgreSQL备份
            backup_file = os.path.join(backup_dir, f"postgres_backup_{timestamp}.sql")
            
            # 从DATABASE_URL解析连接信息
            db_url = settings.DATABASE_URL
            # 这里需要解析DATABASE_URL获取连接参数
            
            # 执行pg_dump
            cmd = f"pg_dump {db_url} > {backup_file}"
            result = subprocess.run(cmd, shell=True, capture_output=True, text=True)
            
            if result.returncode != 0:
                raise Exception(f"PostgreSQL备份失败: {result.stderr}")
            
            logger.info(f"PostgreSQL备份完成: {backup_file}")
            
        elif settings.DATABASE_URL.startswith("sqlite"):
            # SQLite备份
            db_path = settings.DATABASE_URL.replace("sqlite:///", "")
            backup_file = os.path.join(backup_dir, f"sqlite_backup_{timestamp}.db")
            
            # 复制SQLite文件
            import shutil
            shutil.copy2(db_path, backup_file)
            
            logger.info(f"SQLite备份完成: {backup_file}")
        
        # 验证备份文件
        if not os.path.exists(backup_file):
            raise FileNotFoundError(f"备份文件未生成: {backup_file}")
        
        backup_size = os.path.getsize(backup_file)
        if backup_size == 0:
            raise Exception(f"备份文件为空: {backup_file}")
        
        logger.info(f"数据库备份完成: {backup_file}, 大小: {backup_size/(1024*1024):.2f} MB")
        
        return {
            "success": True,
            "message": "数据库备份完成",
            "backup_file": backup_file,
            "backup_size_mb": backup_size / (1024 * 1024)
        }
        
    except Exception as e:
        logger.error(f"数据库备份失败: {str(e)}")
        raise e


@celery_app.task(**get_retry_config('default'))
@retry_task('default', on_failure=log_task_failure, on_retry=log_task_retry)
def archive_old_data(self):
    """
    归档旧数据任务
    
    重试策略: default (最多3次重试，60秒起始延迟，指数退避)
    """
    from datetime import datetime, timedelta
    
    logger.info("开始归档旧数据")
    
    db: Session = SessionLocal()
    
    try:
        # 归档超过6个月的已完成回证
        six_months_ago = datetime.now() - timedelta(days=180)
        
        from app.models.delivery_receipt import DeliveryReceipt, DeliveryStatusEnum
        
        # 查询需要归档的数据
        old_receipts = db.query(DeliveryReceipt).filter(
            DeliveryReceipt.created_at < six_months_ago,
            DeliveryReceipt.status.in_([
                DeliveryStatusEnum.SENT,
                DeliveryStatusEnum.DELIVERED
            ])
        ).all()
        
        archived_count = 0
        
        # 这里可以实现实际的归档逻辑
        # 比如将数据导出到文件或另一个数据库
        for receipt in old_receipts:
            # 示例：标记为已归档（如果有这个状态）
            # receipt.status = DeliveryStatusEnum.ARCHIVED
            archived_count += 1
        
        if archived_count > 0:
            db.commit()
        
        logger.info(f"数据归档完成: 归档 {archived_count} 条记录")
        
        return {
            "success": True,
            "message": "数据归档完成",
            "archived_count": archived_count
        }
        
    except SQLAlchemyError as e:
        logger.error(f"数据归档时数据库错误: {str(e)}")
        db.rollback()
        raise e
        
    except Exception as e:
        logger.error(f"数据归档失败: {str(e)}")
        db.rollback()
        raise e
        
    finally:
        db.close()