#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
insert_imgs_and_title.py
向“送达回证”模板表格
  1. 居中写入“送达文书 名称及文号”（仿宋_GB2312，12 pt）
  2. 插入两张图片并居中

依赖: python-docx  pillow
"""

from pathlib import Path
import argparse
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from PIL import Image


# ---------- 工具函数 ----------

def add_centered_picture(paragraph, img_path: Path, max_width_inch: float):
    """向段落插入等比缩放图片，并让段落居中"""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    with Image.open(img_path) as im:
        w_px, h_px = im.size
        dpi = im.info.get("dpi", (96, 96))[0]

    w_in = w_px / dpi
    h_in = h_px / dpi
    if w_in > max_width_inch:
        scale = max_width_inch / w_in
        w_in *= scale
        h_in *= scale

    run.add_picture(str(img_path), width=Inches(w_in), height=Inches(h_in))


def write_centered_text(cell, text: str):
    """清空 cell 并写入居中文本（仿宋_GB2312，12 pt）"""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = "仿宋_GB2312"
    # 设置 EastAsia 字体，确保中文正确显示
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
    run.font.size = Pt(12)


# ---------- 主流程 ----------

def process_document(
    template_doc: Path,
    output_doc: Path,
    doc_title: str,
    pic_note: Path,
    pic_footer: Path,
    max_width_inch: float = 2.8,
):
    doc = Document(template_doc)
    if not doc.tables:
        raise ValueError("模板中未发现表格")

    table = doc.tables[0]

    # --- 1️⃣  “送达文书 名称及文号” ---
    pos_dispatch = None
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "送达文书" in cell.text:
                pos_dispatch = (r_idx, c_idx)
                break
        if pos_dispatch:
            break
    if not pos_dispatch:
        raise RuntimeError("未找到包含“送达文书”的单元格")

    row_d, col_d = pos_dispatch
    dispatch_target = table.cell(row_d, col_d + 1)  # 右侧单元格
    write_centered_text(dispatch_target, doc_title)

    # --- 2️⃣  “备注”右侧插图 ---
    pos_note = None
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "备注" == cell.text.strip():
                pos_note = (r_idx, c_idx)
                break
        if pos_note:
            break
    if not pos_note:
        raise RuntimeError("未找到包含“备注”的单元格")

    note_target = table.cell(pos_note[0], pos_note[1] + 1)
    note_target.text = ""
    add_centered_picture(note_target.paragraphs[0], pic_note, max_width_inch)

    # --- 3️⃣  最底栏插图 ---
    footer_cell = table.rows[-1].cells[0]
    footer_cell.text = ""
    add_centered_picture(footer_cell.paragraphs[0], pic_footer, max_width_inch)

    # --- 4️⃣  保存 ---
    doc.save(output_doc)
    print(f"✓ 已生成：{output_doc}")


# ---------- 命令行入口 ----------

def main():
    parser = argparse.ArgumentParser(
        description="向“送达回证”模板写入文书标题并插入两张图片"
    )
    parser.add_argument("--template", required=True, help="模板 .docx 路径")
    parser.add_argument("--output", required=True, help="输出 .docx 路径")
    parser.add_argument("--doc-title", required=True, help="送达文书 名称及文号")
    parser.add_argument("--note-img", required=True, help="插入备注栏右侧的图片")
    parser.add_argument("--footer-img", required=True, help="插入最底栏的图片")
    parser.add_argument("--max-width", type=float, default=2.8, help="图片最大宽度 (inch)")

    args = parser.parse_args()

    process_document(
        template_doc=Path(args.template),
        output_doc=Path(args.output),
        doc_title=args.doc_title,
        pic_note=Path(args.note_img),
        pic_footer=Path(args.footer_img),
        max_width_inch=args.max_width,
    )


if __name__ == "__main__":
    main()
